#!/usr/bin/env python3
"""用 OpenAI 兼容小模型验证 BD 检查点执行说明书 + 1 份待审文件。"""

from __future__ import annotations

import argparse
import concurrent.futures
import difflib
import glob
import json
import os
import re
import subprocess
import textwrap
import time
from pathlib import Path
from typing import Any


WORKSPACE_ROOT = Path.cwd()
DEFAULT_OUTPUT_ROOT = WORKSPACE_ROOT / "validation" / "cli-runs"
DEFAULT_JOBS = 8
DEFAULT_MAX_TOKENS = 6144
DEFAULT_MIN_CANDIDATE_SCORE = 3
DEFAULT_CANDIDATE_POOL_FACTOR = 4
DEFAULT_WINDOW_DEDUPE_THRESHOLD = 0.92
DEFAULT_WINDOW_MERGE_GAP = 2
DEFAULT_SHORT_LINE_LOOKAHEAD = 5
DEFAULT_SUPPORT_CONTEXT_WINDOWS = 2
ENV_BASE_URL = "CHECKPOINT_LLM_BASE_URL"
ENV_MODEL = "CHECKPOINT_LLM_MODEL"
ENV_API_KEY = "CHECKPOINT_LLM_API_KEY"


def read_text(path: Path) -> str:
    return path.read_text(encoding="utf-8")


def write_text(path: Path, text: str) -> None:
    path.parent.mkdir(parents=True, exist_ok=True)
    path.write_text(text, encoding="utf-8")


def now_text() -> str:
    return time.strftime("%Y-%m-%d %H:%M:%S")


def run_id() -> str:
    return time.strftime("%Y%m%d-%H%M%S")


def slugify_filename(name: str) -> str:
    sanitized = re.sub(r'[\\/:*?"<>|]+', "-", name).strip()
    return sanitized or "unknown"


def relative_path(path: Path) -> str:
    try:
        return os.path.relpath(path, WORKSPACE_ROOT)
    except Exception:
        return str(path)


def display_file_name(value: Any) -> str:
    text = str(value or "").strip()
    return Path(text).name if text else ""


def under_path(path: Path, parent: Path) -> bool:
    try:
        path.resolve().relative_to(parent.resolve())
        return True
    except Exception:
        return False


def validate_output_dir(output_dir: Path) -> None:
    forbidden = WORKSPACE_ROOT / "wiki" / "audits"
    if under_path(output_dir, forbidden):
        raise RuntimeError("checkpoint CLI 输出不得写入 wiki/audits，请使用 validation/cli-runs。")


def resolve_llm_config(args: argparse.Namespace) -> None:
    args.base_url = args.base_url or os.environ.get(ENV_BASE_URL)
    args.model = args.model or os.environ.get(ENV_MODEL)
    args.api_key = args.api_key or os.environ.get(ENV_API_KEY)
    missing = []
    if not args.base_url:
        missing.append(ENV_BASE_URL)
    if not args.model:
        missing.append(ENV_MODEL)
    if not args.api_key:
        missing.append(ENV_API_KEY)
    if missing:
        raise RuntimeError("缺少模型连接配置，请通过环境变量或命令行参数提供：" + "、".join(missing))


def extract_text_from_docx(path: Path) -> tuple[str, str]:
    textutil = subprocess.run(
        ["textutil", "-convert", "txt", "-stdout", str(path)],
        capture_output=True,
        text=True,
        check=False,
    )
    if textutil.returncode == 0 and textutil.stdout.strip():
        return textutil.stdout, "textutil"

    try:
        import docx  # type: ignore
    except ImportError as exc:
        raise RuntimeError(f"无法提取 {path.name}，textutil 失败且 python-docx 不可用") from exc

    document = docx.Document(str(path))
    lines = [para.text for para in document.paragraphs]
    text = "\n".join(lines).strip()
    if text:
        return text, "python-docx"
    raise RuntimeError(f"无法提取 {path.name}，docx 内容为空")


def normalize_extracted_text(text: str) -> str:
    return (
        text.replace("\r", "")
        .replace("\x0b", "\n")
        .replace("\x0c", "\n")
        .replace("\u2028", "\n")
        .replace("\u2029", "\n")
    )


def extract_structured_review_from_docx(path: Path) -> dict[str, Any]:
    try:
        import docx  # type: ignore
        from docx.document import Document as _Document  # type: ignore
        from docx.oxml.table import CT_Tbl  # type: ignore
        from docx.oxml.text.paragraph import CT_P  # type: ignore
        from docx.table import Table  # type: ignore
        from docx.text.paragraph import Paragraph  # type: ignore
    except ImportError as exc:
        raise RuntimeError(f"无法结构化提取 {path.name}，python-docx 不可用") from exc

    def iter_block_items(parent: _Document):
        for child in parent.element.body.iterchildren():
            if isinstance(child, CT_P):
                yield Paragraph(child, parent)
            elif isinstance(child, CT_Tbl):
                yield Table(child, parent)

    def paragraph_to_text(item: Any) -> str:
        return normalize_extracted_text(str(item.text or "")).strip()

    def table_to_text(item: Any) -> str:
        rows: list[str] = []
        for row in item.rows:
            cells = [normalize_extracted_text(str(cell.text or "")).strip() for cell in row.cells]
            if any(cells):
                rows.append("\t".join(cells))
        return "\n".join(rows).strip()

    document = docx.Document(str(path))
    blocks: list[dict[str, Any]] = []
    all_lines: list[str] = []
    stats = {"paragraph_blocks": 0, "table_blocks": 0, "block_count": 0, "line_count": 0}
    line_no = 1
    block_index = 1
    for item in iter_block_items(document):
        block_type = "paragraph"
        if item.__class__.__name__ == "Paragraph":
            text = paragraph_to_text(item)
            if text:
                stats["paragraph_blocks"] += 1
        else:
            block_type = "table"
            text = table_to_text(item)
            if text:
                stats["table_blocks"] += 1
        if text:
            block_lines = text.split("\n")
            start_line = line_no
            for block_line in block_lines:
                all_lines.append(block_line)
                line_no += 1
            blocks.append(
                {
                    "block_id": f"b{block_index:04d}",
                    "block_type": block_type,
                    "order_index": block_index,
                    "text": text,
                    "lines": block_lines,
                    "line_start": start_line,
                    "line_end": line_no - 1,
                }
            )
            block_index += 1
    stats["block_count"] = len(blocks)
    stats["line_count"] = len(all_lines)
    text = "\n".join(all_lines).strip()
    if not text:
        raise RuntimeError(f"无法结构化提取 {path.name}，docx 内容为空")
    return {
        "text": text,
        "extractor": "python-docx-structured-lines",
        "stats": stats,
        "lines": all_lines,
    }


def extract_text_from_doc(path: Path) -> tuple[str, str]:
    antiword = subprocess.run(
        ["antiword", str(path)],
        capture_output=True,
        text=True,
        check=False,
    )
    if antiword.returncode == 0 and antiword.stdout.strip():
        return antiword.stdout, "antiword"

    textutil = subprocess.run(
        ["textutil", "-convert", "txt", "-stdout", str(path)],
        capture_output=True,
        text=True,
        check=False,
    )
    if textutil.returncode == 0 and textutil.stdout.strip():
        return textutil.stdout, "textutil"

    error_text = antiword.stderr.strip() or textutil.stderr.strip() or "未知错误"
    raise RuntimeError(f"无法提取 {path.name}：{error_text}")


def extract_text_from_pdf(path: Path) -> tuple[str, str]:
    pdftotext = subprocess.run(
        ["pdftotext", "-layout", str(path), "-"],
        capture_output=True,
        text=True,
        check=False,
    )
    if pdftotext.returncode == 0 and pdftotext.stdout.strip():
        return pdftotext.stdout, "pdftotext-layout"

    try:
        import pypdf  # type: ignore
    except ImportError as exc:
        error_text = pdftotext.stderr.strip() or "pdftotext 无输出且 pypdf 不可用"
        raise RuntimeError(f"无法提取 {path.name}：{error_text}") from exc

    reader = pypdf.PdfReader(str(path))
    pages: list[str] = []
    for page_index, page in enumerate(reader.pages, start=1):
        text = normalize_extracted_text(page.extract_text() or "").strip()
        if text:
            pages.append(f"[第 {page_index} 页]\n{text}")
    output = "\n\n".join(pages).strip()
    if output:
        return output, "pypdf"
    error_text = pdftotext.stderr.strip() or "PDF 内容为空或无法抽取文本"
    raise RuntimeError(f"无法提取 {path.name}：{error_text}")


def load_review_file(path: Path) -> tuple[str, str]:
    suffix = path.suffix.lower()
    if suffix in {".md", ".txt"}:
        return read_text(path), "plain-text"
    if suffix == ".docx":
        return extract_text_from_docx(path)
    if suffix == ".doc":
        return extract_text_from_doc(path)
    if suffix == ".pdf":
        return extract_text_from_pdf(path)
    raise RuntimeError(f"暂不支持的待审文件类型：{path.suffix}")


def load_review_file_variants(path: Path) -> dict[str, Any]:
    plain_text, plain_extractor = load_review_file(path)
    variants: dict[str, Any] = {
        "plain_text": plain_text,
        "plain_extractor": plain_extractor,
    }
    if path.suffix.lower() == ".docx":
        try:
            structured_review = extract_structured_review_from_docx(path)
        except Exception as exc:
            variants["structured_error"] = str(exc)
        else:
            variants["structured_text"] = structured_review["text"]
            variants["structured_extractor"] = structured_review["extractor"]
            variants["structured_stats"] = structured_review["stats"]
            variants["structured_lines"] = structured_review["lines"]
    return variants


def parse_frontmatter(markdown: str) -> dict[str, Any]:
    match = re.match(r"(?ms)^---\n(.*?)\n---", markdown)
    if not match:
        return {}
    data: dict[str, Any] = {}
    current_key: str | None = None
    for line in match.group(1).splitlines():
        if not line.strip():
            continue
        key_match = re.match(r"^([A-Za-z0-9_-]+):\s*(.*)$", line)
        if key_match:
            current_key = key_match.group(1)
            value = key_match.group(2).strip()
            data[current_key] = value if value else []
            continue
        list_match = re.match(r"^\s*-\s*(.+)$", line)
        if list_match and current_key:
            value = data.setdefault(current_key, [])
            if isinstance(value, list):
                value.append(list_match.group(1).strip())
    return data


def source_lines(text: str) -> list[str]:
    lines = text.split("\n")
    if lines and lines[-1] == "":
        lines = lines[:-1]
    return lines


def extract_title(markdown: str) -> str:
    frontmatter = parse_frontmatter(markdown)
    title = frontmatter.get("title")
    if isinstance(title, str) and title:
        return title
    match = re.search(r"^#\s+(.+)$", markdown, re.M)
    return match.group(1).strip() if match else "unknown-checkpoint"


def extract_checkpoint_id(markdown: str) -> str:
    frontmatter = parse_frontmatter(markdown)
    checkpoint_id = frontmatter.get("id")
    if isinstance(checkpoint_id, str) and checkpoint_id:
        return checkpoint_id
    return "unknown-id"


def extract_section(markdown: str, heading: str) -> str:
    pattern = re.compile(rf"(?ms)^##\s+{re.escape(heading)}\s*$\n(.*?)(?=^##\s+|\Z)")
    match = pattern.search(markdown)
    return match.group(1).strip() if match else ""


def compact_checkpoint_text(markdown: str, max_chars: int) -> str:
    headings = [
        "审查目标",
        "适用范围",
        "标准检查点",
        "检查点定义",
        "审查问题句",
        "定位与召回剖面",
        "机器召回配置",
        "定位关键词",
        "候选召回规则",
        "上下文读取规则",
        "条款属性判断",
        "相关性判断尺子",
        "专用 SOP 边界",
        "基础命中条件",
        "命中条件",
        "排除条件",
        "判断结果分流",
        "正例",
        "反例",
        "审查步骤",
        "结论表达",
        "风险提示",
        "修改建议",
        "审查依据",
        "示例",
        "易误报场景",
        "地区适用说明",
    ]
    chunks: list[str] = []
    included_headings: set[str] = set()
    for heading in headings:
        section = extract_section(markdown, heading)
        if section:
            chunks.append(f"## {heading}\n{section}")
            included_headings.add(heading)
    for match in re.finditer(r"^##\s+(.+?判断尺子)\s*$", markdown, re.M):
        heading = match.group(1).strip()
        if heading in included_headings:
            continue
        section = extract_section(markdown, heading)
        if section:
            chunks.append(f"## {heading}\n{section}")
            included_headings.add(heading)
    text = "\n\n".join(chunks).strip()
    if not text:
        text = markdown
    if len(text) > max_chars:
        return text[:max_chars] + "\n\n[已按 max-checkpoint-chars 截断]"
    return text


def parse_keyword_groups(markdown: str) -> dict[str, list[str]]:
    section = extract_section(markdown, "定位与召回剖面") or extract_section(markdown, "定位关键词")
    groups: dict[str, list[str]] = {}
    current_group = "关键词"
    groups[current_group] = []
    for line in section.splitlines():
        heading = re.match(r"^###\s+(.+?)\s*$", line)
        if heading:
            current_group = heading.group(1).strip()
            groups.setdefault(current_group, [])
            continue
        item = re.match(r"^-\s+(.+?)\s*$", line)
        if item:
            word = item.group(1).strip().rstrip("。；;")
            if word:
                groups.setdefault(current_group, []).append(word)
    return {key: value for key, value in groups.items() if value}


def line_has_any(line: str, words: list[str]) -> list[str]:
    return [word for word in words if word and word in line]


def words_from_named_groups(keyword_groups: dict[str, list[str]], *name_parts: str) -> list[str]:
    words: list[str] = []
    for group_name, group_words in keyword_groups.items():
        if any(part in group_name for part in name_parts):
            words.extend(group_words)
    return words


def has_scoring_weight_structure(line: str) -> bool:
    """识别待审文件中的评分权重结构，用于补充上下文而非判断风险。"""
    compact = re.sub(r"\s+", "", line)
    return bool(
        re.search(r"权重(?:\(%\)|（%）|百分比|比例)?", compact)
        or re.search(r"评标总得分|综合得分|价格分|技术分|商务分", compact)
        or re.search(r"F\d*[=＝].*A\d*", compact)
        or re.search(r"(评分因素|评分项|评分准则|评分标准)", compact)
    )


def normalize_similarity_text(text: str) -> str:
    compact = normalize_extracted_text(text)
    compact = re.sub(r"[0-9０-９]+", "", compact)
    compact = re.sub(r"[A-Za-z]+", "", compact)
    compact = re.sub(r"[（(][一二三四五六七八九十0-9]+[）)]", "", compact)
    compact = re.sub(r"[^\u4e00-\u9fff]+", "", compact)
    return compact


def merge_hit_words(
    left: dict[str, list[str]],
    right: dict[str, list[str]],
) -> dict[str, list[str]]:
    return {
        group: sorted(set(left.get(group, [])) | set(right.get(group, [])))
        for group in set(left) | set(right)
    }


def window_signature(raw_lines: list[str], start: int, end: int) -> str:
    return normalize_similarity_text("\n".join(raw_lines[start:end]))


def is_similar_window(signature: str, seen_signatures: list[str], threshold: float) -> bool:
    if not signature:
        return False
    return any(difflib.SequenceMatcher(None, signature, seen).ratio() >= threshold for seen in seen_signatures)


def hit_words_by_group_name(hits: dict[str, list[str]], *name_parts: str) -> set[str]:
    words: set[str] = set()
    for group_name, group_words in hits.items():
        if any(part in group_name for part in name_parts):
            words.update(group_words)
    return words


def overlaps_any_window(start: int, end: int, windows: list[tuple[int, int]]) -> bool:
    return any(start <= window_end and end >= window_start for window_start, window_end in windows)


def collect_scoring_weight_context(raw_lines: list[str], max_chars: int = 3000) -> str:
    ranges: list[tuple[int, int]] = []
    for idx, line in enumerate(raw_lines):
        if has_scoring_weight_structure(line):
            ranges.append((max(0, idx - 4), min(len(raw_lines), idx + 24)))

    merged: list[tuple[int, int]] = []
    for start, end in sorted(ranges):
        if not merged or start > merged[-1][1]:
            merged.append((start, end))
        else:
            merged[-1] = (merged[-1][0], max(merged[-1][1], end))

    lines: list[str] = []
    for start, end in merged[:4]:
        for line_no in range(start + 1, end + 1):
            lines.append(f"{line_no:04d}: {raw_lines[line_no - 1]}")
    text = "\n".join(lines)
    if len(text) > max_chars:
        text = text[:max_chars] + "\n[评分折算上下文已截断]"
    return text


def score_keyword_hits(
    text: str,
    keyword_groups: dict[str, list[str]],
) -> tuple[int, dict[str, list[str]], int]:
    table_words = (
        keyword_groups.get("表头词", [])
        + words_from_named_groups(keyword_groups, "表头", "章节", "角色")
    )
    object_words = keyword_groups.get("对象词", []) + words_from_named_groups(keyword_groups, "对象")
    limit_words = (
        keyword_groups.get("限制词", [])
        + words_from_named_groups(keyword_groups, "限制", "行为", "门槛", "模式")
    )
    consequence_words = keyword_groups.get("后果词", []) + words_from_named_groups(keyword_groups, "后果")

    hits = {group: line_has_any(text, words) for group, words in keyword_groups.items()}
    total_hits = sum(len(words) for words in hits.values())
    if total_hits == 0:
        return 0, hits, total_hits

    compact_text = re.sub(r"\s+", "", text)
    scoring_consequence_hits = []
    if re.search(r"得\d+(?:\.\d+)?分", compact_text):
        scoring_consequence_hits.append("得N分")
    if re.search(r"不得分", compact_text):
        scoring_consequence_hits.append("不得分")
    if re.search(r"共计\d+(?:\.\d+)?分", compact_text):
        scoring_consequence_hits.append("共计N分")
    if re.search(r"[（(]\d+(?:\.\d+)?分[）)]", compact_text):
        scoring_consequence_hits.append("括号分值")
    if scoring_consequence_hits:
        hits["评分结构词"] = scoring_consequence_hits

    has_table = bool(line_has_any(text, table_words))
    has_object = bool(line_has_any(text, object_words))
    has_limit = bool(line_has_any(text, limit_words))
    has_consequence = bool(line_has_any(text, consequence_words) or scoring_consequence_hits)

    score = total_hits
    if scoring_consequence_hits:
        score += len(scoring_consequence_hits)
    if has_object and has_limit:
        score += 8
    if has_object and has_consequence:
        score += 8
    if has_table and has_object and has_limit:
        score += 8
    if has_table and has_object and has_consequence:
        score += 10
    if has_object and has_limit and has_consequence:
        score += 10
    if object_words and not has_object and total_hits:
        score -= 1
    return score, hits, total_hits


def collect_candidate_windows(
    review_text: str,
    keyword_groups: dict[str, list[str]],
    checkpoint_id: str,
    checkpoint_title: str,
    context_before: int,
    context_after: int,
    max_windows: int,
    max_line_chars: int,
    max_excerpt_chars: int,
    min_candidate_score: int,
) -> tuple[str, int, dict[str, Any]]:
    raw_lines = source_lines(review_text)
    stats: dict[str, Any] = {
        "raw_hit_count": 0,
        "filtered_hit_count": 0,
        "selected_scores": [],
        "max_score": 0,
        "skip_reason": "",
    }
    if not raw_lines:
        stats["skip_reason"] = "待审文件文本为空"
        return "", 0, stats
    scored: list[tuple[int, int, int, dict[str, list[str]]]] = []
    weak_scored: list[tuple[int, int, int, dict[str, list[str]]]] = []
    for idx, line in enumerate(raw_lines):
        lookahead_end = min(len(raw_lines), idx + DEFAULT_SHORT_LINE_LOOKAHEAD + 1)
        scoring_text = "\n".join(raw_lines[idx:lookahead_end])
        score, hits, total_hits = score_keyword_hits(scoring_text, keyword_groups)
        if total_hits == 0:
            continue
        stats["raw_hit_count"] += 1
        start = max(0, idx - context_before)
        end = min(len(raw_lines), idx + context_after + 1)
        start = expand_template_context_start(raw_lines, start, end)
        if score < min_candidate_score:
            weak_scored.append((score, start, end, hits))
            continue

        stats["filtered_hit_count"] += 1
        scored.append((score, start, end, hits))

    if not scored:
        if stats["raw_hit_count"]:
            stats["skip_reason"] = "仅命中弱关键词，未达到候选窗口提交分数"
        else:
            stats["skip_reason"] = "未命中检查点关键词"
        return "", 0, stats

    scored.sort(key=lambda item: (item[0], -item[1]), reverse=True)
    pool_size = max(max_windows * DEFAULT_CANDIDATE_POOL_FACTOR, max_windows)
    candidate_pool = scored[:pool_size]
    candidate_pool.sort(key=lambda item: item[1])
    stats["candidate_pool_size"] = len(candidate_pool)

    merged_pool: list[tuple[int, int, int, dict[str, list[str]]]] = []
    for score, start, end, hits in candidate_pool:
        if not merged_pool or start > merged_pool[-1][2] + DEFAULT_WINDOW_MERGE_GAP:
            merged_pool.append((score, start, end, hits))
            continue
        prev_score, prev_start, prev_end, prev_hits = merged_pool[-1]
        merged_pool[-1] = (
            max(prev_score, score),
            prev_start,
            max(prev_end, end),
            merge_hit_words(prev_hits, hits),
        )

    stats["merged_window_count"] = len(merged_pool)
    selected: list[tuple[int, int, int, dict[str, list[str]]]] = []
    seen_signatures: list[str] = []
    duplicate_skipped = 0
    for item in sorted(merged_pool, key=lambda row: (row[0], -row[1]), reverse=True):
        _, start, end, _ = item
        signature = window_signature(raw_lines, start, end)
        if is_similar_window(signature, seen_signatures, DEFAULT_WINDOW_DEDUPE_THRESHOLD):
            duplicate_skipped += 1
            continue
        selected.append(item)
        if signature:
            seen_signatures.append(signature)
        if len(selected) >= max_windows:
            break

    selected.sort(key=lambda item: item[1])
    stats["deduped_window_count"] = len(selected)
    stats["duplicate_window_skipped"] = duplicate_skipped
    stats["window_dedupe_threshold"] = DEFAULT_WINDOW_DEDUPE_THRESHOLD
    stats["selected_scores"] = [item[0] for item in selected]
    stats["max_score"] = max(stats["selected_scores"] or [0])

    chunks: list[str] = []
    for idx, (score, start, end, hits) in enumerate(selected, start=1):
        hit_parts = []
        for group, words in hits.items():
            if words:
                hit_parts.append(f"{group}=" + "、".join(words))
        chunk_lines = []
        for line_no in range(start + 1, end + 1):
            line = raw_lines[line_no - 1]
            if len(line) > max_line_chars:
                line = line[:max_line_chars] + "……[本行已截断]"
            chunk_lines.append(f"{line_no:04d}: {line}")
        if any(has_scoring_weight_structure(raw_lines[line_no - 1]) for line_no in range(start + 1, end + 1)):
            chunk_lines.append("")
            chunk_lines.append(
                "评分折算提示：若本窗口表头包含“评分因素 / 权重(%) / 评分准则”，"
                "评分因素名称后的数字通常是该评分因素权重；"
                "条款中的“得N分、优得N分、最高得N分、满分N分”通常是该因素内部得分，"
                "必须结合权重折算后再评价实际总分影响。"
            )
        chunks.append(
            f"[候选窗口 {idx}] score={score}; " + "；".join(hit_parts) + "\n" + "\n".join(chunk_lines)
        )

    selected_object_words: set[str] = set()
    selected_ranges = [(start, end) for _, start, end, _ in selected]
    for _, _, _, hits in selected:
        selected_object_words.update(hit_words_by_group_name(hits, "对象"))

    support_windows: list[tuple[int, int, int, dict[str, list[str]]]] = []
    if selected_object_words:
        weak_scored.sort(key=lambda item: (item[0], -item[1]), reverse=True)
        support_signatures: list[str] = []
        support_ranges: list[tuple[int, int]] = []
        for score, start, end, hits in weak_scored:
            if overlaps_any_window(start, end, selected_ranges) or overlaps_any_window(start, end, support_ranges):
                continue
            weak_object_words = hit_words_by_group_name(hits, "对象")
            if not selected_object_words.intersection(weak_object_words):
                continue
            signature = window_signature(raw_lines, start, end)
            if is_similar_window(signature, support_signatures, DEFAULT_WINDOW_DEDUPE_THRESHOLD):
                continue
            support_windows.append((score, start, end, hits))
            support_ranges.append((start, end))
            if signature:
                support_signatures.append(signature)
            if len(support_windows) >= DEFAULT_SUPPORT_CONTEXT_WINDOWS:
                break

    support_windows.sort(key=lambda item: item[1])
    stats["support_window_count"] = len(support_windows)
    for idx, (score, start, end, hits) in enumerate(support_windows, start=1):
        hit_parts = []
        for group, words in hits.items():
            if words:
                hit_parts.append(f"{group}=" + "、".join(words))
        chunk_lines = []
        for line_no in range(start + 1, end + 1):
            line = raw_lines[line_no - 1]
            if len(line) > max_line_chars:
                line = line[:max_line_chars] + "……[本行已截断]"
            chunk_lines.append(f"{line_no:04d}: {line}")
        chunks.append(
            f"[辅助上下文 {idx}] score={score}; " + "；".join(hit_parts) + "\n" + "\n".join(chunk_lines)
        )

    scoring_context = collect_scoring_weight_context(raw_lines)
    if scoring_context:
        # 候选证据窗口必须优先进入 prompt；评分折算上下文只是辅助材料，
        # 不能因为全局评分表过长而挤掉真正的风险证据。
        chunks.append(
            "[评分折算上下文]\n"
            "以下内容用于判断“内部满分”和“实际总分权重”的关系；"
            "审查评分项时必须优先读取。\n"
            + scoring_context,
        )
    excerpt = "\n\n".join(chunks)
    if len(excerpt) > max_excerpt_chars:
        excerpt = excerpt[:max_excerpt_chars] + "\n\n[候选窗口已按 max-review-excerpt-chars 截断]"
    return excerpt, len(selected), stats


def collect_candidate_windows_result(
    review_text: str,
    keyword_groups: dict[str, list[str]],
    checkpoint_id: str,
    checkpoint_title: str,
    context_before: int,
    context_after: int,
    max_windows: int,
    max_line_chars: int,
    max_excerpt_chars: int,
    min_candidate_score: int,
) -> dict[str, Any]:
    excerpt, window_count, stats = collect_candidate_windows(
        review_text,
        keyword_groups,
        checkpoint_id,
        checkpoint_title,
        context_before,
        context_after,
        max_windows,
        max_line_chars,
        max_excerpt_chars,
        min_candidate_score,
    )
    return {
        "review_excerpt": excerpt,
        "window_count": window_count,
        "recall_stats": stats,
    }


def choose_review_recall(
    review_variants: dict[str, Any],
    keyword_groups: dict[str, list[str]],
    checkpoint_id: str,
    checkpoint_title: str,
    args: argparse.Namespace,
) -> dict[str, Any]:
    review_text = review_variants.get("structured_text") or review_variants["plain_text"]
    extractor = review_variants.get("structured_extractor") or review_variants["plain_extractor"]
    result = collect_candidate_windows_result(
        review_text,
        keyword_groups,
        checkpoint_id,
        checkpoint_title,
        args.context_before,
        args.context_after,
        args.max_windows,
        args.max_line_chars,
        args.max_review_excerpt_chars,
        args.min_candidate_score,
    )
    result["extractor"] = extractor
    result["channel"] = "structured-line" if review_variants.get("structured_text") else "plain-line"
    result["review_text"] = review_text
    result["structured_stats"] = review_variants.get("structured_stats", {})
    result["recall_config"] = {
        "recall_unit": "line",
        "context_before": args.context_before,
        "context_after": args.context_after,
        "max_windows": args.max_windows,
        "max_review_excerpt_chars": args.max_review_excerpt_chars,
        "candidate_pool_factor": DEFAULT_CANDIDATE_POOL_FACTOR,
        "window_dedupe_threshold": DEFAULT_WINDOW_DEDUPE_THRESHOLD,
        "window_merge_gap": DEFAULT_WINDOW_MERGE_GAP,
        "short_line_lookahead": DEFAULT_SHORT_LINE_LOOKAHEAD,
        "support_context_windows": DEFAULT_SUPPORT_CONTEXT_WINDOWS,
    }
    result["fallback_used"] = False
    result["recall_fallback_reason"] = review_variants.get("structured_error", "")
    return result


def build_messages(
    checkpoint_id: str,
    checkpoint_title: str,
    checkpoint_text: str,
    review_name: str,
    review_excerpt: str,
) -> list[dict[str, str]]:
    system = (
        "/no_think\n"
        "你是政府采购招标文件合规审查小模型。"
        "你只能根据给定的 BD 检查点执行说明书和待审文件候选窗口进行审查。"
        "你必须按检查点中的流程执行：候选召回、上下文读取、条款属性判断、相关性三问、基础命中条件、排除条件、核心条件计数、结果分流。"
        "你不能发明未出现在原文中的证据，不能只因关键词命中就输出命中。"
        "不要输出思考过程，不要输出解释性前言，输出必须是 JSON。"
    )
    user = textwrap.dedent(
        f"""
        目标检查点：{checkpoint_id} {checkpoint_title}
        待审文件：{review_name}

        执行要求：
        1. 只能根据下方 BD 检查点执行说明书审查。
        2. 必须先列出候选条款，再逐条执行三问、A/B/C 基础命中、排除条件和核心条件计数。
        3. 最终结论只能是：命中 / 待人工复核 / 不命中。
        4. 如果证据不足，必须输出待人工复核，不能强行命中。
        5. 证据摘录必须来自待审文件候选窗口，不得改写原文。
        6. 必须返回严格合法 JSON：所有字符串内部禁止出现未转义的英文双引号，不要在字符串中用英文双引号列举关键词。
        7. 如果没有候选条款，candidates 必须返回空数组 []，candidate_count 填 0，各步骤 summary 控制在 80 个汉字以内。
        8. 不要复制大段关键词清单，不要输出制表符、异常空白、重复占位符或超长空白。
        9. 审查评分项时，必须区分“评分因素内部得分”和“折算后的评标总分影响”。
        10. 如果原文存在“评标总得分=F×A”或“权重(%)”，不得把“得N分、优得N分、最高得N分、满分N分”直接表述为总分高或分值占比高；必须先读取对应权重。
        11. 无法确认权重时，只能表述为“该评分因素内部得分为 N 分，实际总分影响需结合权重折算”，不得使用“分值高达、分值极高、完全决定评标结果”等表达。
        12. 审查合同或商务条款留白时，必须先判断该条款是正式专用条款，还是“合同条款及格式、合同模板、合同范本、示范文本、仅供参考”类文本。
        13. 如果留白位于仅供参考的合同格式或模板范文中，且采购需求、专用条款或其他正式章节可能补足，不得直接输出命中；应输出待人工复核，并说明需核对正式专用条款或最终签署合同。

        输出 JSON 结构：
        {{
          "checkpoint_id": "{checkpoint_id}",
          "checkpoint_title": "{checkpoint_title}",
          "verdict": "命中|待人工复核|不命中",
          "summary": "一句话总结审查结论",
          "execution_trace": {{
            "candidate_recall": {{"status": "已执行", "summary": "...", "candidate_count": 0}},
            "context_reading": {{"status": "已执行", "summary": "..."}},
            "clause_classification": {{"status": "已执行", "summary": "...", "clause_types": []}},
            "relevance_three_questions": {{"status": "已执行", "summary": "..."}},
            "basic_hit_abc": {{"status": "已执行", "A": false, "B": false, "C": false, "summary": "..."}},
            "exclusion_checks": {{"status": "已执行", "triggered": [], "not_triggered": []}},
            "core_condition_count": {{"status": "已执行", "count": 0, "summary": "..."}},
            "result_branch": {{"status": "已执行", "branch": "命中|待人工复核|不命中", "reason": "..."}}
          }},
          "candidates": [
            {{
              "line_anchor": "行号范围，例如 0123-0128",
              "excerpt": "原文证据摘录",
              "matched_keyword_groups": {{"表头词": [], "对象词": [], "限制词": [], "后果词": []}},
              "clause_type": "评分项|资格条件|符合性审查|实质性响应|证明材料|履约能力说明|模板残留|其他",
              "three_questions": {{
                "certificate_capability": "...",
                "same_capability_in_demand": "是|否|不确定",
                "direct_effect_on_performance": "是|否|不确定"
              }},
              "basic_hit_abc": {{"A": false, "B": false, "C": false}},
              "triggered_exclusions": [],
              "core_conditions_met": [],
              "core_condition_count": 0,
              "candidate_verdict": "命中|待人工复核|不命中",
              "reason": "..."
            }}
          ]
        }}

        【BD 检查点执行说明书】
        {checkpoint_text}

        【待审文件候选窗口】
        {review_excerpt}
        """
    ).strip()
    return [{"role": "system", "content": system}, {"role": "user", "content": user}]


def post_openai_compatible(
    base_url: str,
    api_key: str,
    model: str,
    messages: list[dict[str, str]],
    temperature: float,
    timeout_seconds: int,
    max_tokens: int,
) -> dict[str, Any]:
    url = base_url.rstrip("/") + "/chat/completions"
    payload = {
        "model": model,
        "messages": messages,
        "temperature": temperature,
        "max_tokens": max_tokens,
        "response_format": {"type": "json_object"},
        "chat_template_kwargs": {"enable_thinking": False},
        "enable_thinking": False,
    }
    command = [
        "curl",
        "-sS",
        "--max-time",
        str(timeout_seconds),
        "-H",
        "Content-Type: application/json",
        "-H",
        f"Authorization: Bearer {api_key}",
        "--data",
        json.dumps(payload, ensure_ascii=False),
        url,
    ]
    result = subprocess.run(command, capture_output=True, text=True, check=False)
    if result.returncode != 0:
        raise RuntimeError(f"curl failed: {result.stderr.strip()}")
    try:
        return json.loads(result.stdout)
    except json.JSONDecodeError as exc:
        raise RuntimeError(f"接口未返回合法 JSON：{result.stdout[:1000]}") from exc


def parse_model_json(response: dict[str, Any]) -> dict[str, Any]:
    try:
        content = response["choices"][0]["message"]["content"]
    except (KeyError, IndexError) as exc:
        raise RuntimeError(f"模型响应结构异常：{response}") from exc
    if content is None:
        message = response.get("choices", [{}])[0].get("message", {})
        reasoning = message.get("reasoning")
        if isinstance(reasoning, str) and reasoning.strip().startswith("{"):
            content = reasoning
        else:
            raise RuntimeError("模型返回 content 为空；可能仍在输出 reasoning，请降低 max_tokens 或关闭 thinking")
    try:
        return json.loads(content)
    except json.JSONDecodeError as exc:
        recovered = recover_partial_model_json(content)
        if recovered:
            return recovered
        raise RuntimeError(f"模型未返回合法 JSON：{content[:1000]}") from exc


def recover_partial_model_json(content: str) -> dict[str, Any] | None:
    """Best-effort recovery for small-model responses truncated after key fields."""
    verdict_match = re.search(r'"verdict"\s*:\s*"(命中|待人工复核|不命中)"', content)
    if not verdict_match:
        return None
    def field(name: str) -> str:
        match = re.search(rf'"{re.escape(name)}"\s*:\s*"([^"\\]*(?:\\.[^"\\]*)*)"', content)
        if not match:
            return ""
        try:
            return json.loads(f'"{match.group(1)}"')
        except json.JSONDecodeError:
            return match.group(1)
    return {
        "nbd_id": field("nbd_id") or field("checkpoint_id"),
        "nbd_title": field("nbd_title") or field("checkpoint_title"),
        "checkpoint_id": field("checkpoint_id") or field("nbd_id"),
        "checkpoint_title": field("checkpoint_title") or field("nbd_title"),
        "risk_level": field("risk_level"),
        "result_type": field("result_type") or "其他",
        "verdict": verdict_match.group(1),
        "summary": field("summary") or "模型返回 JSON 截断，已根据 verdict 字段恢复。",
        "candidate_count": 0,
        "execution_trace": {
            "candidate_recall": {"status": "已执行", "summary": "partial-json-recovered"},
            "context_reading": {"status": "已执行", "summary": ""},
            "clause_classification": {"status": "已执行", "summary": "", "clause_types": []},
            "hit_conditions": {"status": "已执行", "A": False, "B": False, "C": False, "summary": ""},
            "exclusion_checks": {"status": "已执行", "triggered": [], "not_triggered": []},
            "result_branch": {"status": "已执行", "branch": verdict_match.group(1), "reason": "partial-json-recovered"},
        },
        "candidates": [],
        "risk_tip": "",
        "revision_suggestion": "",
        "legal_basis": [],
        "recovered_from_invalid_json": True,
    }


def normalize_result(result: dict[str, Any], checkpoint_id: str, checkpoint_title: str) -> dict[str, Any]:
    result.setdefault("checkpoint_id", checkpoint_id)
    result.setdefault("checkpoint_title", checkpoint_title)
    if result.get("verdict") not in {"命中", "待人工复核", "不命中"}:
        result["verdict"] = "待人工复核"
    candidates = result.get("candidates")
    if not isinstance(candidates, list):
        result["candidates"] = []
    trace = result.get("execution_trace")
    if not isinstance(trace, dict):
        trace = {}
    required_steps = [
        "candidate_recall",
        "context_reading",
        "clause_classification",
        "relevance_three_questions",
        "basic_hit_abc",
        "exclusion_checks",
        "core_condition_count",
        "result_branch",
    ]
    for step in required_steps:
        current = trace.get(step)
        if not isinstance(current, dict):
            current = {}
        current.setdefault("status", "已执行")
        current.setdefault("summary", "")
        trace[step] = current
    result["execution_trace"] = trace
    return result


def trace_step_summary(key: str, step: dict[str, Any]) -> str:
    summary = str(step.get("summary", "")).strip()
    if summary:
        return summary
    if key == "exclusion_checks":
        triggered = step.get("triggered")
        not_triggered = step.get("not_triggered")
        parts: list[str] = []
        if isinstance(triggered, list) and triggered:
            parts.append("触发：" + "；".join(str(item) for item in triggered))
        else:
            parts.append("未触发排除条件")
        if isinstance(not_triggered, list) and not_triggered:
            parts.append("未触发项：" + "；".join(str(item) for item in not_triggered))
        return "；".join(parts)
    if key == "result_branch":
        branch = str(step.get("branch", "")).strip()
        reason = str(step.get("reason", "")).strip()
        if branch and reason:
            return f"{branch}：{reason}"
        return branch or reason
    if key == "basic_hit_abc":
        flags = []
        for name in ["A", "B", "C"]:
            if name in step:
                flags.append(f"{name}={step.get(name)}")
        return "，".join(flags)
    return ""


def markdown_report(report: dict[str, Any]) -> str:
    result = report["model_result"]
    candidates = result.get("candidates", [])
    recall_fallback_reason = str(report.get("recall_fallback_reason", "")).strip()
    recall_channel = str(report.get("recall_channel", "")).strip()
    lines = [
        f"# {report['checkpoint_id']} {report['checkpoint_title']} 验证报告",
        "",
        "## 运行信息",
        f"- 开始时间：{report['started_at']}",
        f"- 结束时间：{report['ended_at']}",
        f"- 模型：{report['model']}",
        f"- 检查点：{report['checkpoint_path']}",
        f"- 待审文件：{display_file_name(report.get('review_file', ''))}",
        f"- 文本抽取方式：{report['text_extractor']}",
        f"- 候选召回通道：{recall_channel or 'plain'}",
        f"- 候选窗口数：{report.get('candidate_window_count', 0)}",
        f"- 召回统计：`{json.dumps(report.get('recall_stats', {}), ensure_ascii=False)}`",
        f"- 召回配置：`{json.dumps(report.get('recall_config', {}), ensure_ascii=False)}`",
        "",
        "## 结论",
        f"- 结果：{result.get('verdict', '待人工复核')}",
        f"- 摘要：{result.get('summary', '')}",
        "",
        "## 执行过程",
    ]
    if recall_fallback_reason:
        lines.extend(["", "## 召回说明", f"- {recall_fallback_reason}", ""])
    trace = result.get("execution_trace", {})
    for key, label in [
        ("candidate_recall", "候选召回"),
        ("context_reading", "上下文读取"),
        ("clause_classification", "条款属性判断"),
        ("relevance_three_questions", "相关性三问"),
        ("basic_hit_abc", "基础命中 A/B/C"),
        ("exclusion_checks", "排除条件"),
        ("core_condition_count", "核心条件计数"),
        ("result_branch", "结果分流"),
    ]:
        step = trace.get(key, {})
        lines.append(f"- {label}：{trace_step_summary(key, step)}")
    lines.extend(["", "## 候选条款"])
    if not candidates:
        lines.append("- 未召回候选条款。")
    for idx, candidate in enumerate(candidates, start=1):
        lines.extend(
            [
                f"### 候选 {idx}",
                f"- 行号：{candidate.get('line_anchor', '')}",
                f"- 条款属性：{candidate.get('clause_type', '')}",
                f"- 候选结论：{candidate.get('candidate_verdict', '')}",
                f"- 理由：{candidate.get('reason', '')}",
                "",
                "证据摘录：",
                "",
                "```text",
                str(candidate.get("excerpt", "")).strip(),
                "```",
                "",
            ]
        )
    lines.extend(
        [
            "## 机器可读结果",
            "",
            "```json",
            json.dumps(result, ensure_ascii=False, indent=2),
            "```",
            "",
        ]
    )
    return "\n".join(lines)


def summary_markdown(report: dict[str, Any]) -> str:
    result = report["model_result"]
    return "\n".join(
        [
            f"# {report['checkpoint_id']} 验证摘要",
            "",
            f"- 检查点：{report['checkpoint_title']}",
            f"- 待审文件：{display_file_name(report.get('review_file', ''))}",
            f"- 结果：{result.get('verdict', '待人工复核')}",
            f"- 摘要：{result.get('summary', '')}",
            f"- 召回通道：{report.get('recall_channel', 'plain')}",
            f"- 开始时间：{report['started_at']}",
            f"- 结束时间：{report['ended_at']}",
            f"- 报告：[[{Path(report['report_file']).name}]]",
            "",
        ]
    )


def verdict_rank(verdict: str) -> int:
    return {"命中": 0, "待人工复核": 1, "不命中": 2}.get(verdict, 3)


def load_batch_results(run_dir: Path) -> list[dict[str, Any]]:
    results: list[dict[str, Any]] = []
    for result_file in sorted(run_dir.glob("**/result.json")):
        try:
            data = json.loads(read_text(result_file))
        except json.JSONDecodeError as exc:
            raise RuntimeError(f"无法读取结果文件：{relative_path(result_file)}") from exc
        data["_result_file"] = relative_path(result_file)
        results.append(data)
    if not results:
        raise RuntimeError(f"未在批次目录中找到 result.json：{relative_path(run_dir)}")
    results.sort(
        key=lambda item: (
            verdict_rank(item.get("model_result", {}).get("verdict", "")),
            item.get("checkpoint_id", ""),
        )
    )
    return results


def candidate_summary(candidate: dict[str, Any]) -> str:
    reason = str(candidate.get("reason", "")).strip()
    if reason:
        return reason
    verdict = str(candidate.get("candidate_verdict", "")).strip()
    return f"候选条款结论为{verdict}" if verdict else "模型未提供候选理由。"


def format_candidate_evidence(candidate: dict[str, Any]) -> list[str]:
    lines = [
        f"- 行号证据：`{candidate.get('line_anchor', '')}`",
        f"- 条款属性：{candidate.get('clause_type', '')}",
        f"- 候选结论：{candidate.get('candidate_verdict', '')}",
        f"- 风险原因：{candidate_summary(candidate)}",
    ]
    three_questions = candidate.get("three_questions")
    if isinstance(three_questions, dict):
        lines.extend(
            [
                "- 相关性三问：",
                f"  - 证明能力：{three_questions.get('certificate_capability', '')}",
                f"  - 需求中是否存在同一能力：{three_questions.get('same_capability_in_demand', '')}",
                f"  - 是否直接影响履约：{three_questions.get('direct_effect_on_performance', '')}",
            ]
        )
    basic_hit = candidate.get("basic_hit_abc")
    if isinstance(basic_hit, dict):
        lines.append(
            "- 基础命中 A/B/C："
            f"A={basic_hit.get('A')}, B={basic_hit.get('B')}, C={basic_hit.get('C')}"
        )
    exclusions = candidate.get("triggered_exclusions")
    if isinstance(exclusions, list) and exclusions:
        lines.append("- 触发排除条件：" + "；".join(str(item) for item in exclusions))
    core_conditions = candidate.get("core_conditions_met")
    if isinstance(core_conditions, list) and core_conditions:
        lines.append("- 已满足核心条件：" + "；".join(str(item) for item in core_conditions))
    excerpt = str(candidate.get("excerpt", "")).strip()
    if excerpt:
        lines.extend(["- 证据摘录：", "", "```text", excerpt, "```"])
    return lines


def batch_audit_report_markdown(run_dir: Path, results: list[dict[str, Any]]) -> str:
    started_values = [item.get("started_at", "") for item in results if item.get("started_at")]
    ended_values = [item.get("ended_at", "") for item in results if item.get("ended_at")]
    review_files = sorted({item.get("review_file", "") for item in results if item.get("review_file")})
    models = sorted({item.get("model", "") for item in results if item.get("model")})
    hit_items = [item for item in results if item.get("model_result", {}).get("verdict") == "命中"]
    review_items = [item for item in results if item.get("model_result", {}).get("verdict") == "待人工复核"]
    clean_items = [item for item in results if item.get("model_result", {}).get("verdict") == "不命中"]

    title_review_file = Path(review_files[0]).name if review_files else "未知待审文件"
    review_file_display = display_file_name(review_files[0]) if review_files else ""
    lines = [
        "---",
        f"title: {title_review_file} 小模型检查点审查报告",
        "page_type: checkpoint-cli-audit-report",
        f"run_dir: {relative_path(run_dir)}",
        f"generated_at: {now_text()}",
        "---",
        "",
        f"# {title_review_file} 小模型检查点审查报告",
        "",
        "## 运行信息",
        f"- 批次目录：{relative_path(run_dir)}",
        f"- 待审文件：{review_file_display}",
        f"- 模型：{', '.join(models)}",
        f"- 开始时间：{min(started_values) if started_values else ''}",
        f"- 结束时间：{max(ended_values) if ended_values else ''}",
        f"- 检查点数量：{len(results)}",
        f"- 命中数量：{len(hit_items)}",
        f"- 待人工复核数量：{len(review_items)}",
        f"- 不命中数量：{len(clean_items)}",
        "",
        "## 主审查结论",
    ]
    if hit_items:
        hit_titles = "、".join(f"{item.get('checkpoint_id')} {item.get('checkpoint_title')}" for item in hit_items)
        lines.append(f"本次小模型按检查点 SOP 执行后，认定存在 {len(hit_items)} 项命中风险：{hit_titles}。")
    else:
        lines.append("本次小模型按检查点 SOP 执行后，未输出明确命中风险。")
    if review_items:
        review_titles = "、".join(f"{item.get('checkpoint_id')} {item.get('checkpoint_title')}" for item in review_items)
        lines.append(f"另有 {len(review_items)} 项输出为待人工复核：{review_titles}。这些事项已召回候选证据，但模型认为相关性、必要性、替代路径或排除条件仍需人工确认。")
    if clean_items:
        lines.append(f"其余 {len(clean_items)} 项暂未发现命中情形，列入未命中事项备查。")
    lines.extend(
        [
            "",
            "说明：本报告是业务系统调用检查点 SOP 后的小模型审查产物，作用类似 `full-risk-scan` 的审查结果页；正式定性仍应结合原文、法规依据和人工复核。",
            "",
        ]
    )

    issue_no = 1
    if hit_items:
        lines.append("## 命中风险点")
        for item in hit_items:
            result = item["model_result"]
            candidates = [
                candidate
                for candidate in result.get("candidates", [])
                if candidate.get("candidate_verdict") in {"命中", "待人工复核"}
            ]
            lines.extend(
                [
                    "",
                    f"## 风险点 {issue_no}：{item.get('checkpoint_title')}（{result.get('verdict')}）",
                    "",
                    "- 风险等级：高",
                    f"- 检查点：`{item.get('checkpoint_id')} {item.get('checkpoint_title')}`",
                    f"- 检查点文件：{item.get('checkpoint_path', '')}",
                    f"- 审查结论：{result.get('summary', '')}",
                    "- 项目具体风险：",
                    f"  - {result.get('summary', '')}",
                    "- 风险原因：",
                ]
            )
            if candidates:
                lines.append(f"  - 模型在候选条款中发现 {len(candidates)} 处与该检查点相关的风险或需复核事项。")
            else:
                lines.append("  - 模型未返回可展开候选条款，需回看单点报告。")
            lines.append("- 行号证据：")
            if candidates:
                for candidate in candidates:
                    lines.append(f"  - `{candidate.get('line_anchor', '')}`：{candidate_summary(candidate)}")
            else:
                lines.append("  - 暂无候选行号。")
            lines.extend(["- 候选条款明细：", ""])
            if candidates:
                for candidate_index, candidate in enumerate(candidates, start=1):
                    lines.append(f"### 候选 {candidate_index}")
                    lines.extend(format_candidate_evidence(candidate))
                    lines.append("")
            else:
                lines.append("- 未返回候选条款。")
            lines.extend(
                [
                    "- 回链：",
                    f"  - 单点报告：{item.get('report_file', '')}",
                    f"  - 机器结果：{item.get('_result_file', '')}",
                ]
            )
            issue_no += 1
        lines.append("")

    if review_items:
        lines.append("## 待人工复核事项")
        for item in review_items:
            result = item["model_result"]
            candidates = [
                candidate
                for candidate in result.get("candidates", [])
                if candidate.get("candidate_verdict") in {"命中", "待人工复核"}
            ]
            lines.extend(
                [
                    "",
                    f"## 复核点 {issue_no}：{item.get('checkpoint_title')}（{result.get('verdict')}）",
                    "",
                    "- 风险等级：待人工复核",
                    f"- 检查点：`{item.get('checkpoint_id')} {item.get('checkpoint_title')}`",
                    f"- 检查点文件：{item.get('checkpoint_path', '')}",
                    f"- 审查结论：{result.get('summary', '')}",
                    "- 复核原因：",
                    f"  - {result.get('summary', '')}",
                    "- 行号证据：",
                ]
            )
            if candidates:
                for candidate in candidates:
                    lines.append(f"  - `{candidate.get('line_anchor', '')}`：{candidate_summary(candidate)}")
            else:
                lines.append("  - 暂无候选行号。")
            lines.extend(["- 候选条款明细：", ""])
            if candidates:
                for candidate_index, candidate in enumerate(candidates, start=1):
                    lines.append(f"### 候选 {candidate_index}")
                    lines.extend(format_candidate_evidence(candidate))
                    lines.append("")
            else:
                lines.append("- 未返回候选条款。")
            lines.extend(
                [
                    "- 回链：",
                    f"  - 单点报告：{item.get('report_file', '')}",
                    f"  - 机器结果：{item.get('_result_file', '')}",
                ]
            )
            issue_no += 1
        lines.append("")

    lines.append("## 未命中事项")
    if clean_items:
        for item in clean_items:
            result = item["model_result"]
            lines.extend(
                [
                    f"- `{item.get('checkpoint_id')} {item.get('checkpoint_title')}`：{result.get('summary', '')}",
                    f"  - 单点报告：{item.get('report_file', '')}",
                ]
            )
    else:
        lines.append("- 无。")

    lines.extend(["", "## 回链", f"- 批次日志：{relative_path(run_dir / 'batch.log')}", f"- 结果表：{relative_path(run_dir / 'results.tsv')}"])
    return "\n".join(lines) + "\n"


def parse_line_anchor(anchor: str) -> tuple[int | None, int | None]:
    numbers = [int(item) for item in re.findall(r"\d{1,6}", str(anchor))]
    if not numbers:
        return None, None
    return min(numbers), max(numbers)


def normalize_excerpt_key(excerpt: str) -> str:
    text = re.sub(r"\s+", "", str(excerpt))
    text = re.sub(r"^\d{1,6}[:：]", "", text)
    return text[:160]


def risk_level_for_verdict(verdict: str) -> str:
    if verdict == "命中":
        return "高"
    if verdict == "待人工复核":
        return "待人工复核"
    return "低"


def compact_text_value(value: Any, max_chars: int = 260) -> str:
    text = re.sub(r"\s+", " ", str(value or "")).strip()
    return text if len(text) <= max_chars else text[:max_chars] + "……"


def normalize_score_risk_language(value: Any) -> str:
    text = str(value or "")
    text = re.sub(
        r"该评分项分值高达\s*(\d+(?:\.\d+)?)\s*分",
        r"该评分项内部得分为 \1 分，需结合上级权重折算实际总分影响",
        text,
    )
    text = re.sub(
        r"分值高达\s*(\d+(?:\.\d+)?)\s*分",
        r"内部得分为 \1 分，需结合上级权重折算实际总分影响",
        text,
    )
    text = re.sub(
        r"分值极高[（(](\d+(?:\.\d+)?)\s*分[）)]",
        r"内部得分为 \1 分，需结合上级权重折算实际总分影响",
        text,
    )
    text = re.sub(
        r"分值占比极高",
        "需结合上级权重折算实际总分占比",
        text,
    )
    return text


TEMPLATE_REFERENCE_PHRASES = [
    "仅供参考",
    "具体以项目需求及采购结果为准",
    "以最终签订",
    "最终签署",
    "合同模板",
    "合同范本",
    "示范文本",
    "参考文本",
    "参考格式",
]

TEMPLATE_SECTION_PHRASES = [
    "合同条款及格式",
    "合同条款格式",
    "合同格式",
    "合同草案格式",
    "合同文本格式",
]


def has_placeholder_blank(text: str) -> bool:
    if re.search(r"_{3,}|＿{3,}|…{2,}|\.{4,}", text):
        return True
    return bool(re.search(r"后\s{4,}(?:个工作日|日|天)|前\s{4,}(?:个工作日|日|天)|\s{4,}(?:元|%|％|天|日|个工作日)", text))


def expand_template_context_start(raw_lines: list[str], start: int, end: int) -> int:
    """模板留白候选必须带上方章节标题，否则小模型容易把范本误当正式条款。"""
    window_text = "\n".join(raw_lines[start:end])
    if not has_placeholder_blank(window_text):
        return start
    lookback_start = max(0, start - 60)
    for idx in range(start - 1, lookback_start - 1, -1):
        line = raw_lines[idx]
        if any(phrase in line for phrase in TEMPLATE_REFERENCE_PHRASES + TEMPLATE_SECTION_PHRASES):
            return min(max(0, idx - 2), start)
    return start


def resolve_review_file_path(value: str) -> Path | None:
    if not value:
        return None
    path = Path(value)
    if not path.is_absolute():
        path = WORKSPACE_ROOT / path
    try:
        return path.resolve()
    except Exception:
        return path


def load_review_lines_for_results(results: list[dict[str, Any]]) -> list[str]:
    review_files = sorted({str(item.get("review_file", "")) for item in results if item.get("review_file")})
    if len(review_files) != 1:
        return []
    review_path = resolve_review_file_path(review_files[0])
    if not review_path or not review_path.exists():
        return []
    try:
        review_text, _ = load_review_file(review_path)
    except Exception:
        return []
    return source_lines(review_text)


def template_context_for_candidate(candidate: dict[str, Any], review_lines: list[str]) -> dict[str, Any]:
    start, end = parse_line_anchor(str(candidate.get("line_anchor", "")))
    excerpt = str(candidate.get("excerpt", "") or "")
    reason = candidate_summary(candidate)
    if start is None or end is None or not review_lines:
        context_text = "\n".join([excerpt, reason])
    else:
        lookback_start = max(1, start - 80)
        lookahead_end = min(len(review_lines), end + 5)
        context_text = "\n".join(review_lines[lookback_start - 1 : lookahead_end])
    has_reference = any(phrase in context_text for phrase in TEMPLATE_REFERENCE_PHRASES)
    has_section = any(phrase in context_text for phrase in TEMPLATE_SECTION_PHRASES)
    # 只因候选证据自身存在待填空白才按“模板留白”降级，避免同一合同格式章节内
    # 其他非留白条款被上方模板说明误伤。
    has_blank = has_placeholder_blank(excerpt)
    is_template_blank = has_blank and (has_reference or has_section or "模板" in reason or "最终签署" in reason)
    return {
        "is_template_blank": is_template_blank,
        "has_reference": has_reference,
        "has_section": has_section,
        "has_blank": has_blank,
        "note": "该证据位于合同条款及格式、合同模板或仅供参考文本，且存在待填写空白，不能直接认定为正式商务要求缺失。"
        if is_template_blank
        else "",
    }


def context_text_for_candidate(candidate: dict[str, Any], review_lines: list[str], before: int = 80, after: int = 8) -> str:
    start, end = parse_line_anchor(str(candidate.get("line_anchor", "")))
    excerpt = str(candidate.get("excerpt", "") or "")
    if start is None or end is None or not review_lines:
        return excerpt
    lookback_start = max(1, start - before)
    lookahead_end = min(len(review_lines), end + after)
    return "\n".join(review_lines[lookback_start - 1 : lookahead_end])


def should_suppress_business_candidate(candidate: dict[str, Any], reason: str) -> bool:
    """业务报告只展示有可比证据的风险，不把缺少外部资料本身包装成问题。"""
    text = "\n".join([str(candidate.get("excerpt", "")), reason])
    missing_comparator = any(
        phrase in text
        for phrase in [
            "无公告原文",
            "缺少公告原文",
            "缺公告原文",
            "无法比对",
            "无法确认口径一致性",
            "无法执行比对",
        ]
    )
    only_reference = any(phrase in text for phrase in ["详见公告", "详见招标公告", "引用公告", "按本招标文件第一册第一章招标公告"])
    return missing_comparator or only_reference


def is_supporting_evidence_only(candidate: dict[str, Any], reason: str) -> bool:
    text = "\n".join([str(candidate.get("excerpt", "")), reason])
    return any(phrase in text for phrase in ["综合佐证", "辅助佐证", "线索佐证", "不宜单独认定", "不单独认定"])


def effective_business_verdict(verdict: str, candidate_verdict: str, candidate: dict[str, Any], reason: str) -> str:
    if is_supporting_evidence_only(candidate, reason):
        return "待人工复核"
    return verdict


def common_clause_context_for_candidate(candidate: dict[str, Any], review_lines: list[str]) -> dict[str, Any]:
    context_text = context_text_for_candidate(candidate, review_lines)
    excerpt = str(candidate.get("excerpt", "") or "")
    is_common_terms = (
        "通用条款" in context_text
        and any(phrase in context_text for phrase in ["具有普遍性和通用性", "以专用条款为准", "专用条款”和“通用条款"])
    )
    is_generic_standard = any(
        phrase in excerpt
        for phrase in [
            "设计、制造生产标准及行业标准",
            "符合国家强制性标准要求",
            "符合相关行业标准",
            "符合中华人民共和国的设计、制造生产标准及行业标准",
        ]
    )
    is_common_generic_standard = is_common_terms and is_generic_standard
    return {
        "is_common_generic_standard": is_common_generic_standard,
        "is_common_terms": is_common_terms,
        "is_generic_standard": is_generic_standard,
        "note": "该证据位于第二册通用条款，属于货物质量和标准的通用兜底表述；不能单独认定为项目专用执行标准缺失。"
        if is_common_generic_standard
        else "",
    }


def build_business_issues(results: list[dict[str, Any]], review_lines: list[str] | None = None) -> list[dict[str, Any]]:
    review_lines = review_lines or []
    entries: list[dict[str, Any]] = []
    for item in results:
        result = item.get("model_result", {})
        verdict = str(result.get("verdict", ""))
        if verdict not in {"命中", "待人工复核"}:
            continue
        candidates = result.get("candidates", [])
        if not isinstance(candidates, list):
            candidates = []
        relevant_candidates = [
            candidate
            for candidate in candidates
            if isinstance(candidate, dict) and candidate.get("candidate_verdict") in {"命中", "待人工复核"}
        ]
        if not relevant_candidates:
            # 没有候选证据的待复核只保留在检查点报告，不进入证据视角业务报告。
            continue
        for candidate in relevant_candidates:
            start, end = parse_line_anchor(str(candidate.get("line_anchor", "")))
            reason = normalize_score_risk_language(candidate_summary(candidate))
            checkpoint_id = str(item.get("checkpoint_id", ""))
            if should_suppress_business_candidate(candidate, reason):
                continue
            template_context = template_context_for_candidate(candidate, review_lines)
            common_clause_context = common_clause_context_for_candidate(candidate, review_lines)
            raw_candidate_verdict = str(candidate.get("candidate_verdict", ""))
            base_effective_verdict = effective_business_verdict(verdict, raw_candidate_verdict, candidate, reason)
            effective_verdict = (
                "待人工复核"
                if template_context.get("is_template_blank") or common_clause_context.get("is_common_generic_standard")
                else base_effective_verdict
            )
            effective_candidate_verdict = (
                "待人工复核"
                if template_context.get("is_template_blank") or common_clause_context.get("is_common_generic_standard")
                else effective_business_verdict(raw_candidate_verdict, raw_candidate_verdict, candidate, reason)
            )
            if template_context.get("note"):
                reason = (
                    template_context["note"]
                    + "候选条款存在待填写空白，仅作为复核线索；"
                    + "需核对项目专用条款、采购需求书、补充条款或最终签署合同是否已补足。"
                )
            if common_clause_context.get("note"):
                reason = (
                    common_clause_context["note"]
                    + "需核对第一册专用条款、用户需求书、技术参数表或附件是否已经列明本项目适用的具体标准、规范、检测或验收依据。"
                )
            entries.append(
                {
                    "start": start,
                    "end": end,
                    "line_anchor": str(candidate.get("line_anchor", "")).strip(),
                    "excerpt": str(candidate.get("excerpt", "")).strip(),
                    "excerpt_key": normalize_excerpt_key(str(candidate.get("excerpt", ""))),
                    "verdict": effective_verdict,
                    "clause_type": str(candidate.get("clause_type", "")).strip(),
                    "reason": reason,
                    "template_context": template_context,
                    "common_clause_context": common_clause_context,
                    "checkpoint_item": {
                        "id": checkpoint_id,
                        "title": item.get("checkpoint_title", ""),
                        "verdict": effective_verdict,
                        "summary": normalize_score_risk_language(result.get("summary", "")),
                        "candidate_verdict": effective_candidate_verdict,
                        "reason": reason,
                        "report_file": item.get("report_file", ""),
                        "result_file": item.get("_result_file", ""),
                    },
                }
            )

    groups: list[dict[str, Any]] = []
    anchored = [entry for entry in entries if entry["start"] is not None and entry["end"] is not None]
    unanchored = [entry for entry in entries if entry["start"] is None or entry["end"] is None]
    anchored.sort(key=lambda entry: (entry["start"], entry["end"]))

    def new_group(entry: dict[str, Any]) -> dict[str, Any]:
        return {
            "start": entry["start"],
            "end": entry["end"],
            "line_anchor": entry["line_anchor"],
            "excerpt": entry["excerpt"],
            "verdicts": {entry["verdict"]},
            "checkpoint_items": [entry["checkpoint_item"]],
            "candidate_reasons": [entry["reason"]] if entry["reason"] else [],
            "clause_types": {entry["clause_type"]} if entry["clause_type"] else set(),
            "excerpt_keys": {entry["excerpt_key"]} if entry["excerpt_key"] else set(),
            "template_contexts": [entry["template_context"]] if entry.get("template_context") else [],
            "common_clause_contexts": [entry["common_clause_context"]] if entry.get("common_clause_context") else [],
        }

    for entry in anchored:
        if not groups:
            groups.append(new_group(entry))
            continue
        last = groups[-1]
        same_neighborhood = entry["start"] <= (last["end"] or entry["start"]) + 1
        same_excerpt = bool(entry["excerpt_key"] and entry["excerpt_key"] in last.get("excerpt_keys", set()))
        if same_neighborhood or same_excerpt:
            last["start"] = min(int(last["start"]), int(entry["start"]))
            last["end"] = max(int(last["end"]), int(entry["end"]))
            last["line_anchor"] = f"{int(last['start']):04d}-{int(last['end']):04d}"
            if len(entry["excerpt"]) > len(str(last.get("excerpt", ""))):
                last["excerpt"] = entry["excerpt"]
            last["verdicts"].add(entry["verdict"])
            last["checkpoint_items"].append(entry["checkpoint_item"])
            if entry["reason"]:
                last["candidate_reasons"].append(entry["reason"])
            if entry["clause_type"]:
                last["clause_types"].add(entry["clause_type"])
            if entry["excerpt_key"]:
                last["excerpt_keys"].add(entry["excerpt_key"])
            if entry.get("template_context"):
                last.setdefault("template_contexts", []).append(entry["template_context"])
            if entry.get("common_clause_context"):
                last.setdefault("common_clause_contexts", []).append(entry["common_clause_context"])
        else:
            groups.append(new_group(entry))

    by_excerpt: dict[str, dict[str, Any]] = {}
    for entry in unanchored:
        key = entry["excerpt_key"] or f"unanchored:{len(by_excerpt)}"
        if key not in by_excerpt:
            by_excerpt[key] = new_group(entry)
            continue
        group = by_excerpt[key]
        group["verdicts"].add(entry["verdict"])
        group["checkpoint_items"].append(entry["checkpoint_item"])
        if entry["reason"]:
            group["candidate_reasons"].append(entry["reason"])
        if entry["clause_type"]:
            group["clause_types"].add(entry["clause_type"])
        if entry.get("template_context"):
            group.setdefault("template_contexts", []).append(entry["template_context"])
        if entry.get("common_clause_context"):
            group.setdefault("common_clause_contexts", []).append(entry["common_clause_context"])
    groups.extend(by_excerpt.values())

    issues: list[dict[str, Any]] = []
    sorted_groups = sorted(
        groups,
        key=lambda group: (
            0 if "命中" in group["verdicts"] else 1,
            parse_line_anchor(group.get("line_anchor", ""))[0] or 10**9,
            group.get("line_anchor", ""),
        ),
    )
    for index, group in enumerate(sorted_groups, start=1):
        checkpoint_items_by_key: dict[tuple[str, str], dict[str, Any]] = {}
        for checkpoint_item in group["checkpoint_items"]:
            key = (
                str(checkpoint_item.get("id", "")),
                str(checkpoint_item.get("title", "")),
            )
            existing = checkpoint_items_by_key.get(key)
            if existing is None:
                checkpoint_items_by_key[key] = dict(checkpoint_item)
                continue
            if existing.get("verdict") != "命中" and checkpoint_item.get("verdict") == "命中":
                existing["verdict"] = checkpoint_item.get("verdict", "")
                existing["candidate_verdict"] = checkpoint_item.get("candidate_verdict", "")
            existing_reason = compact_text_value(existing.get("reason", ""), 160)
            new_reason = compact_text_value(checkpoint_item.get("reason", ""), 160)
            if new_reason and new_reason != existing_reason:
                existing["reason"] = "；".join([part for part in [existing_reason, new_reason] if part])
        checkpoint_items = list(checkpoint_items_by_key.values())
        verdict = "命中" if "命中" in group["verdicts"] else "待人工复核"
        template_contexts = group.get("template_contexts", [])
        template_blank = any(context.get("is_template_blank") for context in template_contexts if isinstance(context, dict))
        common_clause_contexts = group.get("common_clause_contexts", [])
        common_generic_standard = any(
            context.get("is_common_generic_standard") for context in common_clause_contexts if isinstance(context, dict)
        )
        first_checkpoint = checkpoint_items[0] if checkpoint_items else {}
        title = str(first_checkpoint.get("title") or "业务风险问题")
        if len(checkpoint_items) > 1:
            title = f"{title}等关联风险"
        reasons = []
        seen_reasons: set[str] = set()
        for reason in group["candidate_reasons"]:
            compact = compact_text_value(reason, 180)
            if compact and compact not in seen_reasons:
                reasons.append(compact)
                seen_reasons.add(compact)
        issues.append(
            {
                "issue_id": f"BI-{index:03d}" if verdict == "命中" else f"BR-{index:03d}",
                "title": title,
                "verdict": verdict,
                "risk_level": risk_level_for_verdict(verdict),
                "line_anchor": group.get("line_anchor", ""),
                "excerpt": group.get("excerpt", ""),
                "clause_types": sorted(group["clause_types"]),
                "triggered_checkpoints": checkpoint_items,
                "risk_analysis": "；".join(reasons[:4]) if reasons else compact_text_value(first_checkpoint.get("summary", "")),
                "manual_review_questions": [
                    "该证据可能位于合同条款及格式、合同模板或仅供参考文本，需先确认其是否属于本项目正式约束条款。",
                    "需核对项目专用条款、采购需求书、补充条款或最终签署合同是否已补足留白事项。",
                ]
                if template_blank
                else [
                    "该证据可能仅是第二册通用条款的兜底表述，需先确认第一册专用条款、用户需求书或技术附件是否已有项目专用标准。",
                    "只有项目专用需求也未列明必要标准、检测或验收依据时，才可进一步认定为采购需求不明确风险。",
                ]
                if common_generic_standard
                else [
                    "需结合采购文件完整上下文、项目实际需求和适用法规进行人工确认。"
                ]
                if verdict == "待人工复核"
                else [],
            }
        )
    return issues


def business_audit_report_data(run_dir: Path, results: list[dict[str, Any]]) -> dict[str, Any]:
    started_values = [item.get("started_at", "") for item in results if item.get("started_at")]
    ended_values = [item.get("ended_at", "") for item in results if item.get("ended_at")]
    review_files = sorted({item.get("review_file", "") for item in results if item.get("review_file")})
    models = sorted({item.get("model", "") for item in results if item.get("model")})
    review_lines = load_review_lines_for_results(results)
    issues = build_business_issues(results, review_lines)
    hit_issues = [issue for issue in issues if issue["verdict"] == "命中"]
    review_issues = [issue for issue in issues if issue["verdict"] == "待人工复核"]
    return {
        "report_type": "business_audit_report",
        "generated_at": now_text(),
        "run_dir": relative_path(run_dir),
        "review_file": review_files[0] if review_files else "",
        "models": models,
        "started_at": min(started_values) if started_values else "",
        "ended_at": max(ended_values) if ended_values else "",
        "checkpoint_count": len(results),
        "summary": {
            "business_issue_count": len(issues),
            "confirmed_issue_count": len(hit_issues),
            "manual_review_count": len(review_issues),
            "checkpoint_hit_count": sum(1 for item in results if item.get("model_result", {}).get("verdict") == "命中"),
            "checkpoint_manual_review_count": sum(1 for item in results if item.get("model_result", {}).get("verdict") == "待人工复核"),
        },
        "issues": issues,
    }


def business_audit_report_markdown(data: dict[str, Any]) -> str:
    review_file_name = display_file_name(data.get("review_file", "")) or "未知待审文件"
    issues = data.get("issues", [])
    hit_issues = [issue for issue in issues if issue.get("verdict") == "命中"]
    review_issues = [issue for issue in issues if issue.get("verdict") == "待人工复核"]
    summary = data.get("summary", {})
    lines = [
        "---",
        f"title: {review_file_name} 业务审查报告",
        "page_type: business-audit-report",
        f"run_dir: {data.get('run_dir', '')}",
        f"generated_at: {data.get('generated_at', '')}",
        "---",
        "",
        f"# {review_file_name} 业务审查报告",
        "",
        "## 一、审查结论摘要",
        "",
        "本报告由 AI 审查生成，用于辅助识别政府采购文件中的合规风险。报告结论不替代采购人、采购代理机构、评审专家、法务人员或监管部门的人工判断。",
        "",
        f"- 审查方式：AI 自动审查",
        f"- 审查模型：{', '.join(data.get('models', []))}",
        f"- 待审文件：{review_file_name}",
        f"- 开始时间：{data.get('started_at', '')}",
        f"- 结束时间：{data.get('ended_at', '')}",
        f"- 检查点数量：{data.get('checkpoint_count', 0)}",
        f"- 形成业务风险问题：{summary.get('confirmed_issue_count', 0)} 个",
        f"- 待人工复核事项：{summary.get('manual_review_count', 0)} 个",
        "",
    ]

    lines.extend(["", "## 二、问题明细", ""])
    if not hit_issues:
        lines.append("本次未形成明确命中的业务风险问题。")
    for index, issue in enumerate(hit_issues, start=1):
        lines.extend(render_business_issue_markdown(issue, f"问题 {index}"))

    lines.extend(["", "## 三、待人工复核事项", ""])
    if not review_issues:
        lines.append("本次未形成按证据聚合的待人工复核事项。")
    for index, issue in enumerate(review_issues, start=1):
        lines.extend(render_business_issue_markdown(issue, f"复核事项 {index}"))

    lines.extend(
        [
            "",
            "## 四、AI审查特别提醒说明",
            "",
            "本报告由 AI 根据标准检查点、待审文件文本和模型审查结果自动生成，主要用于辅助业务人员发现可能存在的风险线索。",
            "",
            "- AI 审查可能存在漏检、误判或对项目背景理解不足的情况。",
            "- 本报告中的“命中”表示 AI 认为存在较明确风险线索，不等同于监管部门的违法认定。",
            "- 本报告中的“待人工复核”表示 AI 发现疑点，但需要结合项目背景、采购需求、公告、更正公告、预算批复、合同附件、交易平台记录等进一步确认。",
            "- 对涉及法律责任、投诉处理、采购文件修改、废标或重新采购等事项，应由业务人员、法务人员或有权监管部门进一步判断。",
            "- 业务人员应结合原始采购文件、项目实际情况和适用法规，对本报告结论进行复核后再使用。",
            "",
        ]
    )
    return "\n".join(lines)


def render_business_issue_markdown(issue: dict[str, Any], label: str) -> list[str]:
    lines = [
        f"### {label}：{issue.get('title', '')}",
        "",
        f"- 风险等级：{issue.get('risk_level', '')}",
        f"- 问题类型：{', '.join(issue.get('clause_types', [])) or '未标注'}",
        f"- 证据位置：{issue.get('line_anchor', '')}",
        f"- 当前状态：{issue.get('verdict', '')}",
        "",
        "#### 原文摘录",
        "",
        "```text",
        str(issue.get("excerpt", "")).strip() or "未返回证据摘录。",
        "```",
        "",
        "#### 触发检查点",
        "",
        "| 检查点 | 结论 | 说明 |",
        "|---|---|---|",
    ]
    for item in issue.get("triggered_checkpoints", []):
        lines.append(
            "| "
            + " | ".join(
                [
                    f"{item.get('id')} {item.get('title')}",
                    str(item.get("verdict", "")),
                    compact_text_value(item.get("reason") or item.get("summary"), 180),
                ]
            )
            + " |"
        )
    lines.extend(
        [
            "",
            "#### 风险说明",
            "",
            issue.get("risk_analysis", "") or "AI 未返回进一步风险说明，需结合单点报告复核。",
            "",
        ]
    )
    if issue.get("manual_review_questions"):
        lines.extend(["#### 需人工确认事项", ""])
        for question in issue.get("manual_review_questions", []):
            lines.append(f"- {question}")
        lines.append("")
    return lines


def write_business_audit_report(run_dir: Path, results: list[dict[str, Any]]) -> tuple[Path, Path]:
    data = business_audit_report_data(run_dir, results)
    report_file = run_dir / "业务审查报告.md"
    data_file = run_dir / "业务审查报告.json"
    write_text(report_file, business_audit_report_markdown(data))
    write_text(data_file, json.dumps(data, ensure_ascii=False, indent=2) + "\n")
    return report_file, data_file


def write_batch_audit_report(run_dir: Path) -> tuple[Path, Path]:
    validate_output_dir(run_dir)
    results = load_batch_results(run_dir)
    report_file = run_dir / "审查报告.md"
    data_file = run_dir / "审查报告.json"
    report_data = {
        "generated_at": now_text(),
        "run_dir": relative_path(run_dir),
        "result_count": len(results),
        "results": results,
    }
    write_text(report_file, batch_audit_report_markdown(run_dir, results))
    write_text(data_file, json.dumps(report_data, ensure_ascii=False, indent=2) + "\n")
    write_business_audit_report(run_dir, results)
    return report_file, data_file


def write_results_tsv(run_dir: Path) -> Path:
    results = load_batch_results(run_dir)
    lines = []
    for item in results:
        result = item.get("model_result", {})
        lines.append(
            "\t".join(
                [
                    str(item.get("checkpoint_id", "")),
                    str(result.get("verdict", "")),
                    str(result.get("summary", "")).replace("\n", " "),
                    str(item.get("_result_file", "")),
                ]
            )
        )
    output_file = run_dir / "results.tsv"
    write_text(output_file, "\n".join(lines) + ("\n" if lines else ""))
    return output_file


def default_output_dir(checkpoint_id: str, review_file: Path) -> Path:
    return DEFAULT_OUTPUT_ROOT / f"checkpoint-{checkpoint_id}-{slugify_filename(review_file.stem)}-{run_id()}"


def default_batch_output_dir(review_file: Path) -> Path:
    return DEFAULT_OUTPUT_ROOT / f"batch-{slugify_filename(review_file.stem)}-{run_id()}"


def default_theme_output_dir(theme_file: Path, review_file: Path) -> Path:
    return DEFAULT_OUTPUT_ROOT / f"theme-{slugify_filename(theme_file.stem)}-{slugify_filename(review_file.stem)}-{run_id()}"


def expand_checkpoint_glob(pattern: str) -> list[Path]:
    matches = [Path(item).resolve() for item in glob.glob(pattern)]
    if not matches:
        matches = [Path(item).resolve() for item in glob.glob(str(WORKSPACE_ROOT / pattern))]
    checkpoint_paths = sorted({path for path in matches if path.is_file() and path.suffix == ".md"})
    if not checkpoint_paths:
        raise RuntimeError(f"未找到检查点文件：{pattern}")
    return checkpoint_paths


def split_markdown_table_row(line: str) -> list[str]:
    stripped = line.strip()
    if stripped.startswith("|"):
        stripped = stripped[1:]
    if stripped.endswith("|"):
        stripped = stripped[:-1]
    cells: list[str] = []
    current: list[str] = []
    wikilink_depth = 0
    idx = 0
    while idx < len(stripped):
        char = stripped[idx]
        pair = stripped[idx : idx + 2]
        if pair == "[[":
            wikilink_depth += 1
            current.append(pair)
            idx += 2
            continue
        if pair == "]]" and wikilink_depth:
            wikilink_depth -= 1
            current.append(pair)
            idx += 2
            continue
        if char == "|" and wikilink_depth == 0:
            cells.append("".join(current).strip())
            current = []
            idx += 1
            continue
        current.append(char)
        idx += 1
    cells.append("".join(current).strip())
    return cells


def extract_bd_ids(value: str) -> list[str]:
    seen: set[str] = set()
    ids: list[str] = []
    for match in re.finditer(r"\bBD\d{2}-\d{3}\b", value):
        bd_id = match.group(0)
        if bd_id not in seen:
            seen.add(bd_id)
            ids.append(bd_id)
    return ids


def resolve_checkpoint_by_id(checkpoint_id: str) -> Path:
    matches = sorted((WORKSPACE_ROOT / "wiki" / "checkpoints").glob(f"{checkpoint_id}*.md"))
    if not matches:
        raise RuntimeError(f"主题引用的检查点不存在：{checkpoint_id}")
    return matches[0].resolve()


def parse_theme_file(theme_file: Path) -> dict[str, Any]:
    theme_text = read_text(theme_file)
    title = extract_title(theme_text) or theme_file.stem
    items: list[dict[str, Any]] = []
    in_mapping = False
    headers: list[str] = []
    for line in theme_text.splitlines():
        if line.startswith("## "):
            in_mapping = "XJGZ 到 BD 组合映射" in line
            continue
        if not in_mapping or not line.strip().startswith("|"):
            continue
        cells = split_markdown_table_row(line)
        if not cells:
            continue
        if cells[0] == "---" or all(set(cell) <= {"-", ":"} for cell in cells if cell):
            continue
        if cells[0] == "业务事项":
            headers = cells
            continue
        if not headers or len(cells) < len(headers):
            continue
        row = dict(zip(headers, cells))
        xjgz_id = row.get("业务事项", "").strip()
        if not xjgz_id:
            continue
        bd_ids = extract_bd_ids(row.get("建议调度 BD", ""))
        items.append(
            {
                "xjgz_id": xjgz_id,
                "description": row.get("检查事项描述", "").strip(),
                "coverage_type": row.get("覆盖类型", "").strip(),
                "bd_ids": bd_ids,
                "coverage_note": row.get("覆盖说明", "").strip(),
            }
        )
    if not items:
        raise RuntimeError(f"主题页未解析到 XJGZ 到 BD 映射表：{relative_path(theme_file)}")
    checkpoint_ids: list[str] = []
    seen: set[str] = set()
    for item in items:
        for bd_id in item["bd_ids"]:
            if bd_id not in seen:
                seen.add(bd_id)
                checkpoint_ids.append(bd_id)
    return {
        "theme_file": relative_path(theme_file),
        "theme_title": title,
        "items": items,
        "checkpoint_ids": checkpoint_ids,
    }


def verdict_rank(verdict: str) -> int:
    return {"命中": 3, "待人工复核": 2, "不命中": 1}.get(verdict, 0)


def theme_item_verdict(results_by_id: dict[str, dict[str, Any]], bd_ids: list[str]) -> str:
    verdict = "不命中"
    for bd_id in bd_ids:
        result = results_by_id.get(bd_id, {}).get("model_result", {})
        current = str(result.get("verdict", ""))
        if verdict_rank(current) > verdict_rank(verdict):
            verdict = current
    return verdict


def write_theme_outputs(theme_dir: Path, theme: dict[str, Any]) -> tuple[Path, Path, Path]:
    results = load_batch_results(theme_dir)
    results_by_id = {str(item.get("checkpoint_id", "")): item for item in results}
    rows: list[dict[str, Any]] = []
    for item in theme["items"]:
        bd_ids = item["bd_ids"]
        matched_bds = [
            bd_id
            for bd_id in bd_ids
            if results_by_id.get(bd_id, {}).get("model_result", {}).get("verdict") in {"命中", "待人工复核"}
        ]
        verdict = theme_item_verdict(results_by_id, bd_ids)
        summaries = [
            compact_text_value(results_by_id.get(bd_id, {}).get("model_result", {}).get("summary", ""), 180)
            for bd_id in matched_bds
        ]
        rows.append(
            {
                "xjgz_id": item["xjgz_id"],
                "description": item["description"],
                "coverage_type": item["coverage_type"],
                "bd_ids": bd_ids,
                "matched_bds": matched_bds,
                "verdict": verdict,
                "summary": "；".join([summary for summary in summaries if summary]) or "未形成命中或待复核 BD 结论。",
                "coverage_note": item["coverage_note"],
            }
        )

    data = {
        "report_type": "theme_validation_report",
        "generated_at": now_text(),
        "run_dir": relative_path(theme_dir),
        "theme": theme,
        "checkpoint_count": len(theme["checkpoint_ids"]),
        "summary": {
            "xjgz_item_count": len(rows),
            "hit_count": sum(1 for row in rows if row["verdict"] == "命中"),
            "manual_review_count": sum(1 for row in rows if row["verdict"] == "待人工复核"),
            "clean_count": sum(1 for row in rows if row["verdict"] == "不命中"),
        },
        "items": rows,
    }
    json_file = theme_dir / "theme.json"
    tsv_file = theme_dir / "theme-results.tsv"
    report_file = theme_dir / "theme-report.md"
    write_text(json_file, json.dumps(data, ensure_ascii=False, indent=2) + "\n")
    tsv_lines = [
        "\t".join(
            [
                row["xjgz_id"],
                row["verdict"],
                row["coverage_type"],
                ",".join(row["bd_ids"]),
                ",".join(row["matched_bds"]),
                row["summary"].replace("\n", " "),
            ]
        )
        for row in rows
    ]
    write_text(tsv_file, "\n".join(tsv_lines) + ("\n" if tsv_lines else ""))

    lines = [
        f"# {theme['theme_title']} 主题验证报告",
        "",
        f"- 主题页：{theme['theme_file']}",
        f"- 运行目录：{relative_path(theme_dir)}",
        f"- XJGZ 事项数：{len(rows)}",
        f"- 调度 BD 数：{len(theme['checkpoint_ids'])}",
        f"- 命中事项：{data['summary']['hit_count']}",
        f"- 待人工复核事项：{data['summary']['manual_review_count']}",
        "",
        "## 主题事项结果",
        "",
        "| XJGZ事项 | 结论 | 覆盖类型 | 调度BD | 命中/复核BD | 摘要 |",
        "|---|---|---|---|---|---|",
    ]
    for row in rows:
        lines.append(
            "| "
            + " | ".join(
                [
                    row["xjgz_id"],
                    row["verdict"],
                    row["coverage_type"],
                    ", ".join(row["bd_ids"]),
                    ", ".join(row["matched_bds"]) or "-",
                    compact_text_value(row["summary"], 220),
                ]
            )
            + " |"
        )
    write_text(report_file, "\n".join(lines) + "\n")
    return json_file, tsv_file, report_file


def build_arg_parser() -> argparse.ArgumentParser:
    parser = argparse.ArgumentParser(
        description="按 BD 检查点 SOP 调用小模型审查 1 份待审文件。"
    )
    parser.add_argument("--checkpoint", type=Path, help="BD 检查点 md 文件路径")
    parser.add_argument("--checkpoint-glob", help="批量验证检查点 glob，例如 'wiki/checkpoints/BD06-*.md'")
    parser.add_argument("--theme-file", type=Path, help="主题组合页 md 文件路径；读取后展开为 BD 检查点批量验证")
    parser.add_argument("--review-file", type=Path, help="待审文件，支持 md/txt/doc/docx")
    parser.add_argument("--aggregate-run-dir", type=Path, help="汇总批次目录下的 result.json，生成审查报告.md")
    parser.add_argument("--output-dir", type=Path, help="输出目录，默认 validation/cli-runs/checkpoint-...")
    parser.add_argument("--base-url", help=f"OpenAI 兼容接口 base_url；也可用环境变量 {ENV_BASE_URL}")
    parser.add_argument("--api-key", help=f"OpenAI 兼容接口密钥；也可用环境变量 {ENV_API_KEY}")
    parser.add_argument("--model", help=f"模型名称；也可用环境变量 {ENV_MODEL}")
    parser.add_argument("--timeout", type=int, default=1800, help="接口超时秒数")
    parser.add_argument("--temperature", type=float, default=0.0, help="模型温度")
    parser.add_argument("--context-before", type=int, default=5, help="候选命中行前文行数")
    parser.add_argument("--context-after", type=int, default=6, help="候选命中行后文行数")
    parser.add_argument("--max-windows", type=int, default=5, help="最多提交候选窗口数")
    parser.add_argument("--max-line-chars", type=int, default=900, help="候选窗口中单行最大字符数")
    parser.add_argument("--max-review-excerpt-chars", type=int, default=8000, help="候选窗口总字符上限")
    parser.add_argument("--max-checkpoint-chars", type=int, default=18000, help="检查点执行说明书字符上限")
    parser.add_argument("--max-tokens", type=int, default=DEFAULT_MAX_TOKENS, help="模型最大输出 token")
    parser.add_argument("--min-candidate-score", type=int, default=DEFAULT_MIN_CANDIDATE_SCORE, help="低于该分数的弱候选不送入模型")
    parser.add_argument("--jobs", type=int, default=DEFAULT_JOBS, help="批量验证并发数")
    parser.add_argument("--resume", action=argparse.BooleanOptionalAction, default=True, help="批量验证时跳过已有 result.json 的检查点")
    parser.add_argument("--reuse-raw-response", action=argparse.BooleanOptionalAction, default=True, help="续跑时复用已存在的 raw-response.json，避免重复请求模型")
    return parser


def run_single_validation(
    args: argparse.Namespace,
    checkpoint_path: Path,
    review_file: Path,
    output_dir_override: Path | None = None,
) -> int:
    checkpoint_path = checkpoint_path.resolve()
    review_file = review_file.resolve()
    checkpoint_text = read_text(checkpoint_path)
    compact_text = compact_checkpoint_text(checkpoint_text, args.max_checkpoint_chars)
    checkpoint_id = extract_checkpoint_id(checkpoint_text)
    checkpoint_title = extract_title(checkpoint_text)
    output_dir = (output_dir_override or args.output_dir or default_output_dir(checkpoint_id, review_file)).resolve()
    validate_output_dir(output_dir)
    output_dir.mkdir(parents=True, exist_ok=True)

    started_at = now_text()
    print(f"start {started_at} checkpoint={checkpoint_id} file={review_file.name}", flush=True)

    try:
        review_variants = load_review_file_variants(review_file)
        keyword_groups = parse_keyword_groups(checkpoint_text)
        recall_choice = choose_review_recall(
            review_variants,
            keyword_groups,
            checkpoint_id,
            checkpoint_title,
            args,
        )
        review_excerpt = str(recall_choice["review_excerpt"])
        window_count = int(recall_choice["window_count"])
        recall_stats = dict(recall_choice["recall_stats"])
        extractor = str(recall_choice["extractor"])
        recall_channel = str(recall_choice.get("channel", "plain"))
        recall_fallback_reason = str(recall_choice.get("recall_fallback_reason", ""))
        fallback_used = bool(recall_choice.get("fallback_used", False))
        prompt_file = output_dir / "prompt.md"
        write_text(
            prompt_file,
            "# Prompt Preview\n\n"
            + f"- checkpoint: {checkpoint_id} {checkpoint_title}\n"
            + f"- review_file: {relative_path(review_file)}\n"
            + f"- recall_channel: {recall_channel}\n"
            + f"- fallback_used: {fallback_used}\n"
            + f"- candidate_windows: {window_count}\n\n"
            + f"- checkpoint_chars: {len(compact_text)}\n"
            + f"- review_excerpt_chars: {len(review_excerpt)}\n\n"
            + f"- recall_stats: {json.dumps(recall_stats, ensure_ascii=False)}\n\n"
            + (
                f"- recall_fallback_reason: {recall_fallback_reason}\n\n"
                if recall_fallback_reason
                else ""
            )
            + (
                f"- structured_stats: {json.dumps(recall_choice.get('structured_stats', {}), ensure_ascii=False)}\n\n"
                if recall_choice.get("structured_stats")
                else ""
            )
            + "## 检查点执行说明书\n\n"
            + compact_text
            + "\n\n"
            + "## 候选窗口\n\n"
            + (review_excerpt or "无有效候选窗口。")
            + "\n",
        )
        print(
            f"recall ok channel={recall_channel} fallback={fallback_used} windows={window_count} "
            f"raw_hits={recall_stats.get('raw_hit_count', 0)} "
            f"filtered_hits={recall_stats.get('filtered_hit_count', 0)} max_score={recall_stats.get('max_score', 0)}",
            flush=True,
        )

        messages = build_messages(
            checkpoint_id,
            checkpoint_title,
            compact_text,
            review_file.name,
            review_excerpt or "无有效候选窗口。",
        )
        raw_response_file = output_dir / "raw-response.json"
        if args.reuse_raw_response and raw_response_file.exists():
            response = json.loads(read_text(raw_response_file))
            model_name = str(response.get("model") or args.model or os.environ.get(ENV_MODEL) or "raw-response")
            print(f"reuse raw-response {relative_path(raw_response_file)}", flush=True)
        else:
            resolve_llm_config(args)
            response = post_openai_compatible(
                args.base_url,
                args.api_key,
                args.model,
                messages,
                args.temperature,
                args.timeout,
                args.max_tokens,
            )
            model_name = str(args.model)
            write_text(raw_response_file, json.dumps(response, ensure_ascii=False, indent=2) + "\n")
        model_result = normalize_result(parse_model_json(response), checkpoint_id, checkpoint_title)
        ended_at = now_text()

        report_file = output_dir / f"{checkpoint_id}-{slugify_filename(checkpoint_title)}.md"
        report: dict[str, Any] = {
            "started_at": started_at,
            "ended_at": ended_at,
            "model": model_name,
            "checkpoint_id": checkpoint_id,
            "checkpoint_title": checkpoint_title,
            "checkpoint_path": relative_path(checkpoint_path),
            "review_file": relative_path(review_file),
            "text_extractor": extractor,
            "recall_channel": recall_channel,
            "recall_fallback_reason": recall_fallback_reason,
            "fallback_used": fallback_used,
            "candidate_window_count": window_count,
            "recall_stats": recall_stats,
            "recall_config": recall_choice.get("recall_config", {}),
            "structured_stats": recall_choice.get("structured_stats", {}),
            "prompt_file": relative_path(prompt_file),
            "report_file": relative_path(report_file),
            "raw_response_file": relative_path(raw_response_file),
            "model_result": model_result,
        }
        write_text(report_file, markdown_report(report))
        write_text(output_dir / "summary.md", summary_markdown(report))
        write_text(output_dir / "result.json", json.dumps(report, ensure_ascii=False, indent=2) + "\n")
        print(f"ok {ended_at} verdict={model_result.get('verdict')} output={relative_path(report_file)}", flush=True)
        return 0
    except Exception as exc:
        ended_at = now_text()
        error_report = {
            "started_at": started_at,
            "ended_at": ended_at,
            "checkpoint_id": checkpoint_id,
            "checkpoint_title": checkpoint_title,
            "checkpoint_path": relative_path(checkpoint_path),
            "review_file": relative_path(review_file),
            "error": str(exc),
        }
        write_text(output_dir / "error.md", "# 运行失败\n\n```json\n" + json.dumps(error_report, ensure_ascii=False, indent=2) + "\n```\n")
        print(f"error {ended_at} {exc}", flush=True)
        return 1


def run_batch_validation(args: argparse.Namespace) -> int:
    if not args.review_file:
        raise SystemExit("批量验证必须提供 --review-file。")
    review_file = args.review_file.resolve()
    checkpoint_paths = expand_checkpoint_glob(args.checkpoint_glob)
    batch_dir = (args.output_dir or default_batch_output_dir(review_file)).resolve()
    validate_output_dir(batch_dir)
    batch_dir.mkdir(parents=True, exist_ok=True)
    batch_log = batch_dir / "batch.log"
    write_text(
        batch_log,
        f"batch_dir={relative_path(batch_dir)}\n"
        f"review_file={relative_path(review_file)}\n"
        f"checkpoint_count={len(checkpoint_paths)}\n"
        f"jobs={args.jobs}\n"
        + f"max_tokens={args.max_tokens}\n"
    )
    failures = 0
    jobs = max(1, int(args.jobs or 1))

    def append_log(text: str) -> None:
        with batch_log.open("a", encoding="utf-8") as handle:
            handle.write(text)

    def run_one(checkpoint_path: Path) -> tuple[str, str, int]:
        checkpoint_text = read_text(checkpoint_path)
        checkpoint_id = extract_checkpoint_id(checkpoint_text)
        checkpoint_title = extract_title(checkpoint_text)
        output_dir = batch_dir / checkpoint_id
        if args.resume and (output_dir / "result.json").exists():
            line = f"\n=== {checkpoint_id} skip existing {now_text()} ===\n"
            append_log(line)
            print(line.strip(), flush=True)
            return checkpoint_id, checkpoint_title, 0
        line = f"\n=== {checkpoint_id} start {now_text()} ===\n"
        append_log(line)
        print(line.strip(), flush=True)
        status = run_single_validation(args, checkpoint_path, review_file, output_dir)
        end_line = f"=== {checkpoint_id} end status={status} {now_text()} ===\n"
        append_log(f"checkpoint_title={checkpoint_title}\n{end_line}")
        print(end_line.strip(), flush=True)
        return checkpoint_id, checkpoint_title, status

    if jobs == 1:
        for checkpoint_path in checkpoint_paths:
            _, _, status = run_one(checkpoint_path)
            if status != 0:
                failures += 1
    else:
        print(f"batch parallel jobs={jobs}", flush=True)
        with concurrent.futures.ThreadPoolExecutor(max_workers=jobs) as executor:
            future_map = {executor.submit(run_one, checkpoint_path): checkpoint_path for checkpoint_path in checkpoint_paths}
            for future in concurrent.futures.as_completed(future_map):
                checkpoint_path = future_map[future]
                try:
                    _, _, status = future.result()
                except Exception as exc:
                    checkpoint_text = read_text(checkpoint_path)
                    checkpoint_id = extract_checkpoint_id(checkpoint_text)
                    append_log(f"=== {checkpoint_id} end status=1 {now_text()} exception={exc} ===\n")
                    print(f"error {checkpoint_id} {exc}", flush=True)
                    status = 1
                if status != 0:
                    failures += 1

    if failures != len(checkpoint_paths):
        results_file = write_results_tsv(batch_dir)
        report_file, data_file = write_batch_audit_report(batch_dir)
        print(
            f"batch report ok results={relative_path(results_file)} output={relative_path(report_file)} data={relative_path(data_file)}",
            flush=True,
        )
    print(f"batch done output={relative_path(batch_dir)} failures={failures}", flush=True)
    return 1 if failures else 0


def run_theme_validation(args: argparse.Namespace) -> int:
    if not args.review_file:
        raise SystemExit("主题验证必须提供 --review-file。")
    theme_file = args.theme_file.resolve()
    review_file = args.review_file.resolve()
    theme = parse_theme_file(theme_file)
    checkpoint_paths = [resolve_checkpoint_by_id(checkpoint_id) for checkpoint_id in theme["checkpoint_ids"]]
    theme_dir = (args.output_dir or default_theme_output_dir(theme_file, review_file)).resolve()
    validate_output_dir(theme_dir)
    theme_dir.mkdir(parents=True, exist_ok=True)
    write_text(theme_dir / "theme-source.json", json.dumps(theme, ensure_ascii=False, indent=2) + "\n")
    batch_log = theme_dir / "batch.log"
    write_text(
        batch_log,
        f"theme_file={theme['theme_file']}\n"
        f"theme_title={theme['theme_title']}\n"
        f"theme_dir={relative_path(theme_dir)}\n"
        f"review_file={relative_path(review_file)}\n"
        f"checkpoint_count={len(checkpoint_paths)}\n"
        f"jobs={args.jobs}\n"
        + f"max_tokens={args.max_tokens}\n",
    )

    failures = 0
    jobs = max(1, int(args.jobs or 1))

    def append_log(text: str) -> None:
        with batch_log.open("a", encoding="utf-8") as handle:
            handle.write(text)

    def run_one(checkpoint_path: Path) -> tuple[str, str, int]:
        checkpoint_text = read_text(checkpoint_path)
        checkpoint_id = extract_checkpoint_id(checkpoint_text)
        checkpoint_title = extract_title(checkpoint_text)
        output_dir = theme_dir / checkpoint_id
        if args.resume and (output_dir / "result.json").exists():
            line = f"\n=== {checkpoint_id} skip existing {now_text()} ===\n"
            append_log(line)
            print(line.strip(), flush=True)
            return checkpoint_id, checkpoint_title, 0
        line = f"\n=== {checkpoint_id} start {now_text()} ===\n"
        append_log(line)
        print(line.strip(), flush=True)
        status = run_single_validation(args, checkpoint_path, review_file, output_dir)
        end_line = f"=== {checkpoint_id} end status={status} {now_text()} ===\n"
        append_log(f"checkpoint_title={checkpoint_title}\n{end_line}")
        print(end_line.strip(), flush=True)
        return checkpoint_id, checkpoint_title, status

    if jobs == 1:
        for checkpoint_path in checkpoint_paths:
            _, _, status = run_one(checkpoint_path)
            if status != 0:
                failures += 1
    else:
        print(f"theme parallel jobs={jobs}", flush=True)
        with concurrent.futures.ThreadPoolExecutor(max_workers=jobs) as executor:
            future_map = {executor.submit(run_one, checkpoint_path): checkpoint_path for checkpoint_path in checkpoint_paths}
            for future in concurrent.futures.as_completed(future_map):
                checkpoint_path = future_map[future]
                try:
                    _, _, status = future.result()
                except Exception as exc:
                    checkpoint_text = read_text(checkpoint_path)
                    checkpoint_id = extract_checkpoint_id(checkpoint_text)
                    append_log(f"=== {checkpoint_id} end status=1 {now_text()} exception={exc} ===\n")
                    print(f"error {checkpoint_id} {exc}", flush=True)
                    status = 1
                if status != 0:
                    failures += 1

    if failures != len(checkpoint_paths):
        results_file = write_results_tsv(theme_dir)
        report_file, data_file = write_batch_audit_report(theme_dir)
        theme_json, theme_tsv, theme_report = write_theme_outputs(theme_dir, theme)
        print(
            f"theme report ok results={relative_path(results_file)} output={relative_path(report_file)} "
            f"theme={relative_path(theme_report)} data={relative_path(theme_json)} tsv={relative_path(theme_tsv)}",
            flush=True,
        )
    print(f"theme done output={relative_path(theme_dir)} failures={failures}", flush=True)
    return 1 if failures else 0


def main() -> int:
    args = build_arg_parser().parse_args()
    if args.aggregate_run_dir:
        run_dir = args.aggregate_run_dir.resolve()
        report_file, data_file = write_batch_audit_report(run_dir)
        print(f"report ok output={relative_path(report_file)} data={relative_path(data_file)}", flush=True)
        return 0
    if args.theme_file:
        return run_theme_validation(args)
    if args.checkpoint_glob:
        return run_batch_validation(args)
    if not args.checkpoint or not args.review_file:
        raise SystemExit("必须提供 --checkpoint 和 --review-file；或使用 --checkpoint-glob 批量验证；或使用 --aggregate-run-dir 汇总批次报告。")
    return run_single_validation(args, args.checkpoint, args.review_file)


if __name__ == "__main__":
    raise SystemExit(main())
